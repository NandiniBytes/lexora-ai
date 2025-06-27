from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from pathlib import Path
from llm_service import query_model
from docx import Document
import time

router = APIRouter()

class NDAGenerateRequest(BaseModel):
    party_1: str
    party_2: str
    jurisdiction: str
    purpose: str

@router.post("/generate")
async def generate_nda(request: NDAGenerateRequest):
    try:
        #loading nda prompt template
        prompt_path = Path("prompts/nda_prompt.txt")
        if not prompt_path.exists():
            raise HTTPException(status_code=500, detail="Prompt template not found")
        
        with prompt_path.open("r") as file:
            prompt_template = file.read()

        #format prompt with request data
        formatted_prompt = prompt_template.format(
            party_1=request.party_1,
            party_2=request.party_2,
            jurisdiction=request.jurisdiction,
            purpose=request.purpose
        )

        #query llm service
        nda_draft = query_model(formatted_prompt)

        #ensure generated_docs directory exists
        docs_dir = Path("generated_docs")
        docs_dir.mkdir(exist_ok=True)

        #generate unique filename with timestamp
        timestamp = int(time.time())
        filename = f"nda_{timestamp}.docx"
        file_path = docs_dir / filename

        #create word document
        doc = Document()
        doc.add_heading("Non-Disclosure Agreement", 0)
        doc.add_paragraph(nda_draft)
        doc.save(file_path)

        return {
            "nda_draft": nda_draft,
            "download_url": f"/api/nda/download/{filename}"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating NDA: {str(e)}")

@router.get("/download/{filename}")
async def download_nda(filename: str):
    file_path = Path("generated_docs") / filename
    if not file_path.exists() or not file_path.is_file():
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )