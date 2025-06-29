from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from database import get_db
from models import Query, Log
from llm_service import query_model
from rag_pipeline import run_rag_pipeline
from pathlib import Path
from docx import Document
import time
import logging

logger = logging.getLogger(__name__)
router = APIRouter()

class NDAGenerateRequest(BaseModel):
    party_1: str
    party_2: str
    jurisdiction: str
    purpose: str

class NDAResponse(BaseModel):
    success: bool
    nda_draft: str
    download_url: str
    rag_enhanced: bool

@router.post("/generate", response_model=NDAResponse)
async def generate_nda(request: NDAGenerateRequest, db: Session = Depends(get_db)):
    try:
        # First, get RAG-enhanced context about NDAs in the specific jurisdiction
        rag_question = f"What are the key requirements for Non-Disclosure Agreements in {request.jurisdiction}? Include standard clauses and legal considerations."
        rag_context = run_rag_pipeline(rag_question)
        
        # Load NDA prompt template
        prompt_path = Path("prompts/nda_prompt.txt")
        if not prompt_path.exists():
            raise HTTPException(status_code=500, detail="Prompt template not found")
        
        with prompt_path.open("r") as file:
            prompt_template = file.read()

        # Enhanced prompt with RAG context
        enhanced_prompt = f"""
        {prompt_template}
        
        Additional Legal Context from Legal Documents:
        {rag_context}
        
        Please ensure the NDA complies with the legal requirements mentioned above.
        """

        # Format prompt with request data
        formatted_prompt = enhanced_prompt.format(
            party_1=request.party_1,
            party_2=request.party_2,
            jurisdiction=request.jurisdiction,
            purpose=request.purpose
        )

        # Query LLM service
        nda_draft = query_model(formatted_prompt)

        # Ensure generated_docs directory exists
        docs_dir = Path("generated_docs")
        docs_dir.mkdir(exist_ok=True)

        # Generate unique filename with timestamp
        timestamp = int(time.time())
        filename = f"nda_{timestamp}.docx"
        file_path = docs_dir / filename

        # Create Word document
        doc = Document()
        doc.add_heading("Non-Disclosure Agreement", 0)
        doc.add_paragraph(nda_draft)
        doc.save(file_path)

        # Save query to database
        query_record = Query(
            question=f"NDA Generation for {request.party_1} and {request.party_2}",
            context=f"Jurisdiction: {request.jurisdiction}, Purpose: {request.purpose}",
            response=nda_draft
        )
        db.add(query_record)
        
        # Log the event
        log_record = Log(
            event_type="nda_generation_rag",
            details=f"Generated RAG-enhanced NDA for {request.party_1} and {request.party_2}"
        )
        db.add(log_record)
        db.commit()

        return NDAResponse(
            success=True,
            nda_draft=nda_draft,
            download_url=f"/api/nda/download/{filename}",
            rag_enhanced=True
        )
        
    except Exception as e:
        logger.error(f"Error generating NDA: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating NDA: {str(e)}")

@router.get("/download/{filename}")
async def download_nda(filename: str):
    file_path = Path("generated_docs") / filename
    if not file_path.exists() or not file_path.is_file():
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
